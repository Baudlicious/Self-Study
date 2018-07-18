#!/usr/bin/env python
# This is x3830s python files from the book Automate Python The
# Boring Stuff with Python. 

def title(big_title, date):
    print('!!!!' + big_title.upper().center(70) + ' !!!!')
    print(date.upper().center(80))

def sep(title):
    print("-" * 80)
    print('[ + ] ' + title.title())
    print("-" * 80 +"\n")

def function(func):
    print(">> " + func + ":")

title('Editing Docx files with Python', '18/07/18')

import docx

sep('Open the document to be edited with docx.Document("docname")')
de = docx.Document('demo.docx')

sep('View paragraphs with .paragraphs()')
print(de.paragraphs)
print(de.paragraphs[0].text)
print(de.paragraphs[1].text)

sep('When there is a change in style such as bold text in a paragraph we can use .runs')
p = de.paragraphs[1]
print(p.runs)
print(p.runs[0].text)
print(p.runs[1].text)
print(p.runs[2].text)

function('We can also see if a condition is true or not for bold or italic with .italic . bold')
print(p.runs[0].bold)
print(p.runs[1].italic)
print(p.runs[2].underline)

function('We can assign bold, italic, and underline to the values')
p.runs[0].bold = True
p.runs[1].italic = True
p.runs[2].underline = True 
print(p.runs[0].bold)
print(p.runs[1].italic)
print(p.runs[2].underline)

function('We can change the text as well by assigning it')
print(p.runs[3].text)

p.runs[3].text = 'italic and underline.'
print(p.runs[3].text)

sep('To save the document we use .save()')
de.save('demo2.docx')

sep('The .style can be used for styles in the doc such as Title')
print(p.style)
p.style = 'Title'
print(p.style)
de.save('demo3.docx')

sep('To create a completely new document use docx.Document()')
d = docx.Document()

sep('To add a paragraph we use .add_paragraph()')
d.add_paragraph('Hello this is a paragraph')
d.add_paragraph('This is another paragraph')
d.save('demo4.docx')

sep('Runs can also be added')
p = d. paragraphs[0]
print(p.runs)
p.add_run('This is a new run.')
print(p.runs)
p.runs[1].bold = True
d.save('demo5.docx')

sep('Get all items in a document as a string')
def get_text(filename):
    doc = docx.Document(filename)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.txt)
    return '\n'.join(full_text)

print(get_text('demo.docx'))


